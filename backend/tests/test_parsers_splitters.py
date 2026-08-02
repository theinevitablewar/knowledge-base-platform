import fitz
import pytest
from docx import Document as WordDocument

from app.rag.parsers import parser_for
from app.rag.splitters import chunker_for


@pytest.mark.asyncio
@pytest.mark.parametrize("extension", [".txt", ".md"])
async def test_text_and_markdown_parsers(tmp_path, extension):
    path = tmp_path / f"sample{extension}"
    path.write_text("# 标题\n\n第一段内容。\n\n第二段内容。", encoding="utf-8-sig")
    parsed = await parser_for(path.name).parse(str(path))
    assert parsed.pages[0].content.startswith("# 标题")
    chunks = chunker_for("markdown" if extension == ".md" else "recursive", 100, 10).split(parsed)
    assert chunks and all(item.content for item in chunks)


@pytest.mark.asyncio
async def test_docx_parser(tmp_path):
    path = tmp_path / "sample.docx"
    document = WordDocument()
    document.add_heading("制度", 1)
    document.add_paragraph("审批流程正文")
    document.save(path)
    parsed = await parser_for(path.name).parse(str(path))
    assert "审批流程正文" in parsed.pages[0].content


@pytest.mark.asyncio
async def test_pdf_parser_preserves_page_number(tmp_path):
    path = tmp_path / "sample.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Approval policy")
    pdf.save(path)
    pdf.close()
    parsed = await parser_for(path.name).parse(str(path))
    assert parsed.pages[0].page_number == 1
    assert "Approval policy" in parsed.pages[0].content
