import asyncio
from pathlib import Path

from docx import Document as DocxDocument

from app.core.exceptions import ParserError
from app.rag.types import ParsedDocument, ParsedPage

from .common import normalize_text


class DocxParser:
    async def parse(self, file_path: str) -> ParsedDocument:
        def parse_sync() -> ParsedDocument:
            try:
                document = DocxDocument(file_path)
                content = normalize_text("\n".join(p.text for p in document.paragraphs))
                title = document.core_properties.title or Path(file_path).stem
                return ParsedDocument(
                    title=title,
                    pages=[ParsedPage(page_number=None, title=title, content=content, metadata={})],
                    metadata={"source_file": Path(file_path).name},
                )
            except Exception as exc:
                raise ParserError("DOCX 解析失败") from exc

        return await asyncio.to_thread(parse_sync)
